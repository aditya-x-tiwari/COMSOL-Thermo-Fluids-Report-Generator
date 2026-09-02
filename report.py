"""
report.py

Builds the Word validation report from a folder of case zips (each
produced by pipeline.py, containing plots/ + manifest.json) plus,
optionally, the source paper's PDF for side-by-side figure comparison.

    python report.py --artifacts-dir artifacts --output report.docx
    python report.py --artifacts-dir artifacts --paper paper.pdf \\
                      --figure-map figure_map.json --output report.docx

figure_map.json (write once per paper - maps case_name -> which paper
figure it validates against):
    {
      "case1_Re100_Ri10": {
        "pdf_page": 7,
        "bbox": [150, 120, 480, 300],
        "label": "Fig. 3, Ri=10, Re=100"
      }
    }
Find pdf_page/bbox once by rendering a page (render_page below) and
opening the PNG in any image viewer to read off pixel coordinates.

Sections:
    1. FIGURE EXTRACTION - crop a paper figure straight from its PDF
    2. DOCX ASSEMBLY     - side-by-side table per case
    3. ARTIFACT LOADING  - unzip case zips, read their manifests
    4. CLI
"""

from __future__ import annotations

import argparse
import json
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


# ======================================================================
# 1. FIGURE EXTRACTION
# ======================================================================

def render_page(pdf_path: str, page_number: int, dpi: int = 200,
                 out_dir: Path = Path("output/pdf_pages")) -> Path:
    """page_number is 1-indexed, the PDF's physical page order (not
    necessarily the printed page label in a journal's header)."""
    out_dir.mkdir(parents=True, exist_ok=True)
    prefix = out_dir / f"page-{page_number}"
    subprocess.run(
        ["pdftoppm", "-png", "-r", str(dpi), "-f", str(page_number), "-l", str(page_number),
         pdf_path, str(prefix)],
        check=True,
    )
    candidates = list(out_dir.glob(f"page-{page_number}*.png"))
    if not candidates:
        raise FileNotFoundError(f"pdftoppm produced nothing for page {page_number}")
    return candidates[0]


def crop_figure(page_png: Path, bbox: tuple, out_path: Path) -> Path:
    """bbox = (left, top, right, bottom) in pixels, at the DPI used to render."""
    img = Image.open(page_png)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.crop(bbox).save(out_path)
    print(f"[report] cropped {bbox} from {page_png.name} -> {out_path}")
    return out_path


def extract_paper_figure(pdf_path: str, page_number: int, bbox: tuple, out_path: str, dpi: int = 200) -> Path:
    return crop_figure(render_page(pdf_path, page_number, dpi), bbox, Path(out_path))


# ======================================================================
# 2. DOCX ASSEMBLY
# ======================================================================

def build_docx(cases: list[dict], output_path: str, title: str = "COMSOL Validation Report"):
    """
    cases: list of dicts, each:
        {"name": str, "description": str,
         "plots": {plot_name: Path}, "paper_figure": {"label": str, "image_file": str} (optional)}
    """
    doc = Document()
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for case in cases:
        doc.add_heading(f"{case['name']} - {case.get('description', '')}", level=1)

        plot_paths = list(case.get("plots", {}).values())
        paper_fig = case.get("paper_figure")
        has_paper_fig = bool(paper_fig) and Path(paper_fig["image_file"]).exists()

        n_cols = len(plot_paths) + (1 if has_paper_fig else 0)
        if n_cols == 0:
            doc.add_paragraph("(no plots for this case)")
            continue

        table = doc.add_table(rows=2, cols=n_cols)
        table.autofit = True
        col = 0
        for img_path in plot_paths:
            cell = table.cell(0, col)
            cell.paragraphs[0].add_run().add_picture(str(img_path), width=Inches(6.0 / n_cols))
            cap = table.cell(1, col).paragraphs[0]
            cap.add_run(f"COMSOL: {Path(img_path).stem}").italic = True
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            col += 1

        if has_paper_fig:
            cell = table.cell(0, col)
            cell.paragraphs[0].add_run().add_picture(paper_fig["image_file"], width=Inches(6.0 / n_cols))
            cap = table.cell(1, col).paragraphs[0]
            cap.add_run(paper_fig.get("label", "Paper figure")).italic = True
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[report] saved -> {output_path}")
    return output_path


# ======================================================================
# 3. ARTIFACT LOADING
# ======================================================================

def load_manifests(artifacts_dir: Path, extract_dir: Path) -> list[dict]:
    manifests = []
    for zip_path in sorted(artifacts_dir.glob("*.zip")):
        case_dir = extract_dir / zip_path.stem
        with zipfile.ZipFile(zip_path) as zf:
            zf.extractall(case_dir)
        manifest_path = case_dir / "manifest.json"
        if not manifest_path.exists():
            print(f"[report] WARNING: {zip_path.name} has no manifest.json, skipping")
            continue
        manifest = json.loads(manifest_path.read_text())
        manifest["_extract_dir"] = case_dir
        manifests.append(manifest)
    return manifests


def build_cases_from_manifests(manifests: list[dict], pdf_path: str | None,
                                figure_map: dict, paper_fig_dir: Path) -> list[dict]:
    cases = []
    for m in manifests:
        name = m["case_name"]
        plots_dir = m["_extract_dir"] / "plots"
        plots = {pname: plots_dir / fname for pname, fname in m["plots"].items()}

        params_str = ", ".join(f"{k}={v:g}" for k, v in m.get("params", {}).items())
        derived_str = ", ".join(f"{k}={v:.4g}" for k, v in m.get("derived", {}).items())
        description = f"{params_str} | {derived_str}" if derived_str else params_str

        case = {"name": name, "description": description, "plots": plots}

        fig_spec = figure_map.get(name)
        if fig_spec and pdf_path:
            out_path = paper_fig_dir / f"{name}_paper.png"
            extract_paper_figure(pdf_path, fig_spec["pdf_page"], tuple(fig_spec["bbox"]), str(out_path))
            case["paper_figure"] = {"label": fig_spec.get("label", "Paper figure"), "image_file": str(out_path)}

        cases.append(case)
    return cases


# ======================================================================
# 4. CLI
# ======================================================================

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--artifacts-dir", default="artifacts")
    parser.add_argument("--paper", default=None, help="Path to the source paper PDF")
    parser.add_argument("--figure-map", default=None, help="JSON: case_name -> {pdf_page, bbox, label}")
    parser.add_argument("--output", default="output/validation_report.docx")
    args = parser.parse_args()

    figure_map = {}
    if args.figure_map and Path(args.figure_map).exists():
        figure_map = json.loads(Path(args.figure_map).read_text())

    with tempfile.TemporaryDirectory() as tmp:
        manifests = load_manifests(Path(args.artifacts_dir), Path(tmp) / "extracted")
        if not manifests:
            raise SystemExit(f"No valid case zips (with manifest.json) found in {args.artifacts_dir}")

        paper_fig_dir = Path("output/paper_figures")
        paper_fig_dir.mkdir(parents=True, exist_ok=True)
        cases = build_cases_from_manifests(manifests, args.paper, figure_map, paper_fig_dir)
        build_docx(cases, args.output)
